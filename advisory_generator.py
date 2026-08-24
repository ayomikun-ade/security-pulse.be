import json
import os
import logging
from io import BytesIO

from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

SYSTEM_PROMPT = """You are a professional cybersecurity analyst writing formal security advisories for IT and security teams.

Given a news article about a cybersecurity threat, produce a structured advisory with exactly these six fields.
Return ONLY a valid JSON object — no markdown, no explanation, no code fences.

{
  "title": "Short professional title, e.g. 'Critical Advisory: ...' or 'Advisory: ...'",
  "overview": "2-3 sentences describing what happened, the threat actor or vulnerability, and current exploitation status.",
  "impact": "4-6 sentences on what a successful compromise enables — data theft, RCE, privilege escalation, etc.",
  "affected_products": "Newline-separated list of affected vendors, products, and versions.",
  "preventive_measures": "Newline-separated list of 4-6 concrete, actionable mitigation steps.",
  "references": "Newline-separated list of relevant URLs (include the source article URL)."
}

Write in formal, precise language. Do not speculate beyond what the article supports."""


def generate_advisory_sections(
    title: str,
    url: str,
    source: str,
    article_text: str,
    description: str = "",
) -> dict:
    body = article_text or description or title
    user_content = (
        f"Article Title: {title}\n"
        f"Source: {source}\n"
        f"URL: {url}\n\n"
        f"Article Content:\n{body[:7000]}"
    )

    response = client.chat.completions.create(
        model="openai/gpt-oss-120b",
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_content},
        ],
        temperature=0.2,
        max_tokens=1500,
        response_format={"type": "json_object"},
    )

    raw = response.choices[0].message.content.strip()

    # Strip accidental markdown fences
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1].lstrip("json").strip() if len(parts) > 1 else raw

    try:
        sections = json.loads(raw)
    except json.JSONDecodeError:
        # Sanitise stray control characters (literal newlines inside JSON strings)
        import re
        cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', raw)
        sections = json.loads(cleaned)

    required = {"title", "overview", "impact", "affected_products", "preventive_measures", "references"}
    missing = required - sections.keys()
    if missing:
        raise ValueError(f"Groq response missing fields: {missing}")

    return sections


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------

_BG_HEADER = "0d1f3c"  # dark navy — header ID cell and section label rows
_BG_WHITE  = "ffffff"  # title cell and content rows
_FG_WHITE  = "ffffff"
_FG_BLUE   = "2563eb"  # title text colour
_FG_BODY   = "1a1a1a"  # body text

_COL_ID    = Inches(1.5)
_COL_TITLE = Inches(5.1)


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _make_run(para, text: str, bold: bool = False, color: str = _FG_BODY):
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Verdana"
    run.font.size = Pt(11)
    r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)


def _merged_row(table):
    row = table.add_row()
    return row.cells[0].merge(row.cells[1])


def _add_section_label(table, label: str):
    cell = _merged_row(table)
    cell.paragraphs[0].clear()
    _make_run(cell.paragraphs[0], label.upper(), bold=True, color=_FG_WHITE)
    _set_cell_bg(cell, _BG_HEADER)


def _add_content_text(table, text: str):
    cell = _merged_row(table)
    cell.paragraphs[0].clear()
    _make_run(cell.paragraphs[0], text.strip(), color=_FG_BODY)
    _set_cell_bg(cell, _BG_WHITE)


def _add_content_list(table, text: str, numbered: bool = False):
    cell = _merged_row(table)
    _set_cell_bg(cell, _BG_WHITE)
    items = [line.strip() for line in text.strip().splitlines() if line.strip()]
    for i, item in enumerate(items):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.clear()
        para.paragraph_format.left_indent = Inches(0.2)
        prefix = f"{i + 1}.  " if numbered else "•  "
        _make_run(para, prefix + item, color=_FG_BODY)


def build_advisory_docx(sections: dict, advisory_id: str) -> bytes:
    doc = Document()

    for sec in doc.sections:
        sec.top_margin = Inches(0.9)
        sec.bottom_margin = Inches(0.9)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)

    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    # ── Header row: ID | Title (2 columns) ─────────────────────────────────
    header_row = table.add_row()
    header_row.cells[0].width = _COL_ID
    header_row.cells[1].width = _COL_TITLE

    id_cell = header_row.cells[0]
    id_cell.paragraphs[0].clear()
    _make_run(id_cell.paragraphs[0], advisory_id, bold=True, color=_FG_WHITE)
    _set_cell_bg(id_cell, _BG_HEADER)

    title_cell = header_row.cells[1]
    title_cell.paragraphs[0].clear()
    _make_run(title_cell.paragraphs[0], sections.get("title", ""), bold=True, color=_FG_BLUE)
    _set_cell_bg(title_cell, _BG_WHITE)

    # ── Section rows (label + content, each full-width merged) ─────────────
    _add_section_label(table, "Overview")
    _add_content_text(table, sections.get("overview", ""))

    _add_section_label(table, "Impact")
    _add_content_text(table, sections.get("impact", ""))

    _add_section_label(table, "Affected Products")
    _add_content_list(table, sections.get("affected_products", ""), numbered=False)

    _add_section_label(table, "Preventive Measures")
    _add_content_list(table, sections.get("preventive_measures", ""), numbered=False)

    _add_section_label(table, "References")
    _add_content_list(table, sections.get("references", ""), numbered=True)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
